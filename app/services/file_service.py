import os
import PyPDF2
import docx
from typing import Dict, List
from datetime import datetime
from io import BytesIO


class FileService:
    """Service for processing uploaded files and extracting text content"""
    
    ALLOWED_EXTENSIONS = {'.txt', '.pdf', '.doc', '.docx'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        pass
    
    def validate_file(self, filename: str, file_size: int) -> Dict:
        """Validate file extension and size"""
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in self.ALLOWED_EXTENSIONS:
            return {
                "valid": False,
                "error": f"File type {ext} not supported. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}"
            }
        
        if file_size > self.MAX_FILE_SIZE:
            return {
                "valid": False,
                "error": f"File too large. Max size: {self.MAX_FILE_SIZE / 1024 / 1024}MB"
            }
        
        return {"valid": True}
    
    def extract_text_from_txt(self, file_content: bytes) -> Dict:
        """Extract text from TXT file"""
        try:
            text = file_content.decode('utf-8')
            return {
                "success": True,
                "text": text,
                "word_count": len(text.split())
            }
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_content.decode(encoding)
                    return {
                        "success": True,
                        "text": text,
                        "word_count": len(text.split())
                    }
                except:
                    continue
            
            return {"success": False, "error": "Could not decode text file"}
    
    def extract_text_from_pdf(self, file_content: bytes) -> Dict:
        """Extract text from PDF file"""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "page_count": len(pdf_reader.pages),
                "word_count": len(full_text.split())
            }
        except Exception as e:
            return {"success": False, "error": f"PDF extraction failed: {str(e)}"}
    
    def extract_text_from_docx(self, file_content: bytes) -> Dict:
        """Extract text from DOCX file"""
        try:
            doc_file = BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "paragraph_count": len(doc.paragraphs),
                "word_count": len(full_text.split())
            }
        except Exception as e:
            return {"success": False, "error": f"DOCX extraction failed: {str(e)}"}
    
    def process_file(self, filename: str, file_content: bytes) -> Dict:
        """Process uploaded file and extract text"""
        # Validate
        validation = self.validate_file(filename, len(file_content))
        if not validation["valid"]:
            return validation
        
        # Extract based on file type
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.txt':
            result = self.extract_text_from_txt(file_content)
        elif ext == '.pdf':
            result = self.extract_text_from_pdf(file_content)
        elif ext in ['.doc', '.docx']:
            result = self.extract_text_from_docx(file_content)
        else:
            return {"success": False, "error": f"Unsupported file type: {ext}"}
        
        if result.get("success"):
            result["filename"] = filename
            result["file_type"] = ext
            result["processed_at"] = datetime.utcnow().isoformat()
        
        return result
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for embedding"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk:
                chunks.append(chunk)
        
        return chunks


# Singleton instance
file_service = FileService()
